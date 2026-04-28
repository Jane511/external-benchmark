"""Shared python-docx primitives for committee reports.

Separated from governance.py / benchmark_report.py so the same styling,
table defaults, and section scaffolding flow through every generated
document. Every helper lazy-imports python-docx — callers get a clear
ImportError pointing at the optional extra if the dependency is missing.

Design choices:
  - Page: US Letter, 1 inch margins all sides.
  - Fonts: Arial (requested). Body 11pt, H2 14pt bold, H1 18pt bold.
  - Tables: explicit column widths, single-pixel borders (via the default
    Table Grid style), header-row bold + light shading.
  - No reliance on custom Word styles beyond "Table Grid" (ships with Word).
"""
from __future__ import annotations

from typing import Any, Iterable, Optional


# ---------------------------------------------------------------------------
# Lazy-import gate
# ---------------------------------------------------------------------------

def _require_docx():
    """Return python-docx's (Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH).

    Raises ImportError with a clear install hint when python-docx is missing.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise ImportError(
            "DOCX export requires python-docx. Install with: "
            "pip install external_benchmark_engine[reports]"
        ) from exc
    return {
        "Document": Document,
        "Inches": Inches,
        "Pt": Pt,
        "RGBColor": RGBColor,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "qn": qn,
        "OxmlElement": OxmlElement,
    }


# ---------------------------------------------------------------------------
# Document bootstrap
# ---------------------------------------------------------------------------

def new_document(title: str, subtitle: Optional[str] = None):
    """Create a Document with US Letter + 1in margins + Arial defaults + a title block."""
    d = _require_docx()
    Document = d["Document"]
    Inches = d["Inches"]
    Pt = d["Pt"]

    doc = Document()

    # Page size + margins on every section.
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default paragraph style -> Arial 11pt.
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    # Title block (H1)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)

    if subtitle:
        sub_para = doc.add_paragraph()
        sub_run = sub_para.add_run(subtitle)
        sub_run.font.name = "Arial"
        sub_run.font.size = Pt(12)
        sub_run.italic = True

    return doc


# ---------------------------------------------------------------------------
# Headings + paragraphs
# ---------------------------------------------------------------------------

def add_heading(doc, text: str, level: int = 2) -> None:
    """Add a styled heading. level=1 -> H1 (18pt), level=2 -> H2 (14pt), else 12pt bold."""
    d = _require_docx()
    Pt = d["Pt"]

    size_map = {1: 18, 2: 14, 3: 12}
    size = size_map.get(level, 12)

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(size)


def add_paragraph(doc, text: str, *, italic: bool = False, bold: bool = False) -> None:
    d = _require_docx()
    Pt = d["Pt"]

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.italic = italic
    run.bold = bold


def add_bullet(doc, text: str) -> None:
    """Bullet via the built-in 'List Bullet' style (ships with Word)."""
    d = _require_docx()
    Pt = d["Pt"]

    try:
        p = doc.add_paragraph(style="List Bullet")
    except KeyError:
        p = doc.add_paragraph()
    run = p.add_run(text) if not p.runs else p.runs[0]
    if not p.runs or p.runs[0].text != text:
        run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def add_table(
    doc,
    *,
    headers: list[str],
    rows: Iterable[list[Any]],
    widths_inches: Optional[list[float]] = None,
    header_shading: str = "DDDDDD",
):
    """Add a bordered table with header styling.

    `headers` is a list of column labels. `rows` is any iterable of lists
    (one per row). `widths_inches` is optional explicit column widths in
    inches — when omitted, Word distributes evenly.

    Returns the docx Table object for further customisation.
    """
    d = _require_docx()
    Inches = d["Inches"]
    Pt = d["Pt"]
    qn = d["qn"]
    OxmlElement = d["OxmlElement"]

    rows_list = list(rows)
    n_cols = len(headers)
    try:
        table = doc.add_table(rows=1 + len(rows_list), cols=n_cols, style="Table Grid")
    except KeyError:
        table = doc.add_table(rows=1 + len(rows_list), cols=n_cols)

    # Explicit column widths if provided
    if widths_inches:
        for col_idx, w in enumerate(widths_inches):
            for row in table.rows:
                if col_idx < len(row.cells):
                    row.cells[col_idx].width = Inches(w)

    # Header row: bold + shading
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        if i >= len(header_cells):
            break
        cell = header_cells[i]
        # Clear any default paragraph and add a fresh styled one.
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)

        # Light-grey shading via XML (python-docx doesn't expose shading directly).
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), header_shading)
        tc_pr.append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows_list, start=1):
        if r_idx >= len(table.rows):
            break
        row_cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row_data):
            if c_idx >= len(row_cells):
                break
            cell = row_cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(_format_cell(val))
            run.font.name = "Arial"
            run.font.size = Pt(11)
    return table


def _format_cell(val: Any) -> str:
    if val is None:
        return ""
    if isinstance(val, float):
        # Common rates like 0.0072 -> "0.72%" is more committee-readable, but
        # many of our values aren't rates. Keep it generic and precise.
        if abs(val) < 1 and val != 0:
            return f"{val:.6f}".rstrip("0").rstrip(".") or "0"
        return str(val)
    return str(val)


# ---------------------------------------------------------------------------
# Footer (institution-agnostic)
# ---------------------------------------------------------------------------

def set_footer(doc, text: str) -> None:
    """Plain footer text. Page numbers require more XML plumbing; skipping for now.

    The spec asks for "page number + 'Generated by External Benchmark Engine'".
    We include the engine tagline; page numbers are a follow-up that needs
    OxmlElement field code surgery (not blocking for V1).
    """
    d = _require_docx()
    Pt = d["Pt"]

    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = ""
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.italic = True


# ---------------------------------------------------------------------------
# Bank-specific block: 3 Lines of Defence sign-off tables (MRC format)
# ---------------------------------------------------------------------------

def add_3lod_signoff(doc) -> None:
    """MRC-format: one sign-off table per line of defence, empty rows to fill in."""
    add_heading(doc, "3 Lines of Defence Sign-Off", level=2)
    add_paragraph(
        doc,
        "Each line of defence signs below to confirm review and acceptance "
        "of this benchmark governance report.",
        italic=True,
    )
    for role, label in [
        ("1st Line of Defence", "Model Owner"),
        ("2nd Line of Defence", "Model Validation"),
        ("3rd Line of Defence", "Internal Audit"),
    ]:
        add_heading(doc, f"{role} — {label}", level=3)
        add_table(
            doc,
            headers=["Name", "Role", "Date", "Signature"],
            rows=[["", label, "", ""]],
            widths_inches=[1.8, 1.8, 1.2, 1.7],
        )


# ---------------------------------------------------------------------------
# Private-credit block: Credit Committee decision log + next actions
# ---------------------------------------------------------------------------

def add_decision_log(doc, *, entries: Optional[list[dict[str, str]]] = None) -> None:
    """Credit-committee format: decision log table.

    When `entries` is empty, render a blank row ready for manual fill-in.
    """
    add_heading(doc, "Decision Log", level=2)
    add_paragraph(
        doc,
        "Record each Credit Committee decision arising from this report.",
        italic=True,
    )
    rows: list[list[str]] = []
    if entries:
        for e in entries:
            rows.append([
                e.get("date", ""),
                e.get("decision", ""),
                e.get("rationale", ""),
                e.get("owner", ""),
                e.get("status", ""),
            ])
    else:
        rows = [["", "", "", "", ""]]
    add_table(
        doc,
        headers=["Date", "Decision", "Rationale", "Owner", "Status"],
        rows=rows,
        widths_inches=[0.9, 1.7, 2.3, 1.0, 0.6],
    )


def add_next_review_actions(doc, *, actions: Optional[list[str]] = None) -> None:
    add_heading(doc, "Next Review Actions", level=2)
    if not actions:
        actions = [
            "[ ] Refresh stale sources before next committee review",
            "[ ] Review any divergence flags raised in the Pillar 3 peer comparison",
            "[ ] Confirm coverage >= 2 sources for every active segment",
        ]
    for item in actions:
        add_bullet(doc, item)
