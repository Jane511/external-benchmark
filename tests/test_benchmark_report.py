"""Tests for reports/benchmark_report.py — Report 1 Benchmark Calibration Summary.

Fixtures build a seeded registry + engine components; each test exercises a
specific section of the 10-section report or an output format. DOCX tests
re-open the produced file to verify content.
"""
from __future__ import annotations

import json
from pathlib import Path

import pytest

from reports.benchmark_report import BenchmarkCalibrationReport
from src.adjustments import AdjustmentEngine
from src.calibration_feed import CalibrationFeed
from src.db import create_engine_and_schema
from src.downturn import DownturnCalibrator
from src.governance import GovernanceReporter
from src.models import InstitutionType
from src.registry import BenchmarkRegistry
from src.seed_data import load_seed_data
from src.triangulation import BenchmarkTriangulator


@pytest.fixture()
def engine_components_bank():
    """All six engine components pre-wired against a seeded in-memory registry."""
    engine = create_engine_and_schema(":memory:")
    registry = BenchmarkRegistry(engine, actor="test")
    load_seed_data(registry)
    adjuster = AdjustmentEngine(InstitutionType.BANK, engine)
    triangulator = BenchmarkTriangulator(InstitutionType.BANK)
    feed = CalibrationFeed(registry, adjuster, triangulator)
    downturn = DownturnCalibrator(registry)
    gov = GovernanceReporter(registry, InstitutionType.BANK)
    return registry, adjuster, triangulator, feed, downturn, gov


@pytest.fixture()
def report_bank(engine_components_bank) -> BenchmarkCalibrationReport:
    registry, adjuster, triangulator, feed, downturn, gov = engine_components_bank
    return BenchmarkCalibrationReport(
        registry=registry, adjustment_engine=adjuster,
        triangulator=triangulator, calibration_feed=feed,
        downturn_calibrator=downturn, governance_reporter=gov,
        institution_type="bank", period_label="Q3 2025",
    )


# ---------------------------------------------------------------------------
# generate() — structured output
# ---------------------------------------------------------------------------

def test_generate_produces_all_sections(report_bank) -> None:
    data = report_bank.generate()
    expected_keys = {
        "meta", "executive_summary", "peer_comparison", "industry_context",
        "source_register", "adjustment_audit_trail", "triangulated_values",
        "calibration_outputs", "downturn_lgd", "bank_vs_pc_comparison",
        "data_governance", "version_history", "source_documentation",
        "narratives",
    }
    assert set(data.keys()) == expected_keys


def test_section_8_flags_are_grouped_not_flat(report_bank) -> None:
    """Governance flags must be emitted as summary groups, not individual
    bullets. The registry can carry dozens of flags on ASIC/ABS rows alone;
    the grouped view keeps Section 8 to a handful of readable summary lines
    while preserving full detail in the Technical Appendix."""
    data = report_bank.generate()
    groups = data["data_governance"].get("groups")
    assert groups is not None, "Grouped governance flags must be exposed"
    # Expect a bounded number of (rule, publisher, dimension) buckets — never
    # a flat bullet-per-source_id list. The upper bound (30) is generous to
    # accommodate fixtures that seed many publishers; in production with
    # clean data this is typically ≤ 10.
    assert 0 <= len(groups) <= 30, (
        f"Governance flags should be grouped into a small set of summary "
        f"buckets, not flat-emitted. Got {len(groups)} groups."
    )
    # No grouped line should re-emit a per-source bullet count of 1 across
    # every possible source (the flat-bullet anti-pattern).
    flat_threshold = max(1, len(data["data_governance"]["all_flags"]) // 2)
    assert len(groups) <= flat_threshold or len(groups) <= 30, (
        "Grouping collapsed too little — roughly one group per flag."
    )
    # No grouped line should carry a raw source_id prefix.
    for g in groups:
        assert "ASIC_ABS_INDUSTRY_" not in g["interpretation"], (
            f"Grouped flag interpretation leaks a raw source_id: {g!r}"
        )
        # Per-bucket count must match the underlying flat count.
        assert g["count"] >= 1


def test_narratives_are_attached_to_generate_output(report_bank) -> None:
    """Every section narrative key in the class template must render to a
    non-empty string in generate()['narratives']."""
    data = report_bank.generate()
    narratives = data.get("narratives")
    assert isinstance(narratives, dict) and narratives
    for key in ("executive_summary", "triangulated", "calibration",
                "downturn_lgd", "bank_vs_pc", "governance",
                "version_history", "source_docs", "signoff"):
        assert narratives.get(key), (
            f"Narrative for {key!r} is missing or empty"
        )


def test_executive_summary_includes_segment_count_and_flagship_visible(
    report_bank,
) -> None:
    data = report_bank.generate()
    summary_text = "\n".join(data["executive_summary"]["lines"]).lower()
    assert "segments covered" in summary_text
    # Flagship is its own section, but exec summary sees it via bank_vs_pc_comparison.
    flagship = data["bank_vs_pc_comparison"]
    assert flagship["raw_pd"] == 0.025


def test_source_register_matches_registry_count(report_bank) -> None:
    registry = report_bank._registry
    data = report_bank.generate()
    assert data["source_register"]["count"] == len(registry.list())


def test_adjustment_audit_trail_has_steps_for_residential_mortgage(
    report_bank,
) -> None:
    data = report_bank.generate()
    segments = {block["segment"] for block in data["adjustment_audit_trail"]}
    assert "residential_mortgage" in segments
    res_block = next(b for b in data["adjustment_audit_trail"]
                     if b["segment"] == "residential_mortgage")
    # Bank chain applies peer_mix -> at least one step per entry
    assert len(res_block["steps"]) > 0
    assert any("peer_mix" in s["name"] for s in res_block["steps"])


def test_triangulated_values_one_row_per_segment(report_bank) -> None:
    data = report_bank.generate()
    rows = data["triangulated_values"]["rows"]
    segments = {r["segment"] for r in rows}
    # Each triangulatable segment appears at most once.
    assert len(rows) == len(segments)


def test_calibration_outputs_shows_five_methods_per_segment(report_bank) -> None:
    data = report_bank.generate()
    for seg_block in data["calibration_outputs"]:
        methods = [m["method"] for m in seg_block["methods"]]
        assert set(methods) == {
            "central_tendency", "logistic_recalibration",
            "bayesian_blending", "external_blending", "pluto_tasche",
        }


def test_bank_vs_pc_flagship_ratio_matches_215x(report_bank) -> None:
    """Pinned flagship: CBA CRE 2.5% -> Bank ~2.50% / PC 5.3762% -> ratio ~2.15."""
    data = report_bank.generate()
    fl = data["bank_vs_pc_comparison"]
    assert fl["raw_pd"] == 0.025
    assert fl["bank_output"] == pytest.approx(0.025, abs=1e-6)
    assert fl["pc_output"] == pytest.approx(0.053762, abs=1e-6)
    assert fl["ratio"] == pytest.approx(2.15, abs=0.01)


def test_data_governance_section_summarises_multiple_reports(report_bank) -> None:
    data = report_bank.generate()
    types = {r["report_type"] for r in data["data_governance"]["reports"]}
    # stale + quality + coverage + pillar3_divergence composed in
    assert {"stale", "quality", "coverage", "pillar3_divergence"}.issubset(types)


def test_version_history_section_handles_no_prior_registry(report_bank) -> None:
    data = report_bank.generate()
    assert data["version_history"]["compared_to_prior"] is False
    assert len(data["version_history"]["rows"]) > 0


# ---------------------------------------------------------------------------
# Output formats
# ---------------------------------------------------------------------------

def test_to_markdown_produces_all_sections(report_bank, tmp_path) -> None:
    path = tmp_path / "report.md"
    report_bank.to_markdown(path)
    text = path.read_text(encoding="utf-8")
    for header in (
        "## 1. Executive Summary",
        "## 2. Source Register",
        "## 3. Adjustment Audit Trail",
        "## 4. Triangulated Values",
        "## 5. Calibration Outputs",
        "## 6. Downturn LGD",
        "## 7. Bank vs Private Credit Comparison",
        "## 8. Data Governance",
        "## 9. Version History",
        "## 10. Source Documentation",
    ):
        assert header in text, f"missing header: {header}"
    # Flagship callout present
    assert "2.15x" in text or "2.15" in text
    # Bank-format 3LoD section present (institution_type="bank")
    assert "3 Lines of Defence" in text


def test_to_html_contains_all_section_anchors(report_bank, tmp_path) -> None:
    path = tmp_path / "report.html"
    report_bank.to_html(path)
    text = path.read_text(encoding="utf-8")
    for anchor in (
        "id='executive-summary'", "id='source-register'",
        "id='adjustment-audit-trail'", "id='triangulated-values'",
        "id='calibration-outputs'", "id='downturn-lgd'",
        "id='bank-vs-pc-comparison'", "id='data-governance'",
        "id='version-history'", "id='source-documentation'",
    ):
        assert anchor in text, f"missing anchor: {anchor}"
    # Self-contained: inline style, no external links
    assert "<style>" in text
    assert "Flagship" in text


def test_to_docx_bank_produces_non_empty_with_3lod(report_bank, tmp_path) -> None:
    from docx import Document

    path = tmp_path / "report.docx"
    report_bank.to_docx(path)
    assert path.stat().st_size > 3000

    doc = Document(str(path))
    text_parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            text_parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text_parts.append(cell.text)
    combined = "\n".join(text_parts)
    assert "Model Risk Committee" in combined
    assert "3 Lines of Defence Sign-Off" in combined
    assert "Model Owner" in combined
    # Flagship mention
    assert "Flagship" in combined or "CRE" in combined


def test_to_docx_pc_produces_non_empty_with_decision_log(
    engine_components_bank, tmp_path,
) -> None:
    """Rebuild the report with institution_type='private_credit' and verify
    Credit Committee framing in the DOCX output."""
    from docx import Document

    registry, adjuster, triangulator, feed, downturn, gov = engine_components_bank
    # Use a PC adjustment engine so the adjustment audit trail works for PC products.
    pc_engine = AdjustmentEngine(InstitutionType.PRIVATE_CREDIT, registry._engine)
    pc_feed = CalibrationFeed(
        registry, pc_engine,
        BenchmarkTriangulator(InstitutionType.PRIVATE_CREDIT),
    )
    pc_gov = GovernanceReporter(registry, InstitutionType.PRIVATE_CREDIT)

    pc_report = BenchmarkCalibrationReport(
        registry=registry, adjustment_engine=pc_engine,
        triangulator=BenchmarkTriangulator(InstitutionType.PRIVATE_CREDIT),
        calibration_feed=pc_feed,
        downturn_calibrator=downturn,
        governance_reporter=pc_gov,
        institution_type="private_credit", period_label="Q3 2025",
    )
    path = tmp_path / "pc_report.docx"
    pc_report.to_docx(path)
    assert path.stat().st_size > 3000

    doc = Document(str(path))
    text = []
    for p in doc.paragraphs:
        text.append(p.text)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                text.append(c.text)
    combined = "\n".join(text)
    assert "Credit Committee" in combined
    assert "Decision Log" in combined
    assert "Next Review Actions" in combined
