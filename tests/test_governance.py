"""Tests for src/governance.py — read-only observer, 6 reports, DOCX shim."""
from __future__ import annotations

from datetime import date
from pathlib import Path
from unittest.mock import patch

import pytest

from src.db import Adjustment, AuditLog, create_engine_and_schema, make_session_factory
from src.governance import (
    GovernanceReporter,
    export_to_docx,
    load_refresh_schedules,
)
from src.models import (
    BenchmarkEntry,
    DataType,
    InstitutionType,
    QualityScore,
    SourceType,
)
from src.registry import BenchmarkRegistry
from src.seed_data import load_seed_data
from sqlalchemy import func, select


def _entry(**overrides) -> BenchmarkEntry:
    base = {
        "source_id": "X", "publisher": "P",
        "source_type": SourceType.PILLAR3, "data_type": DataType.PD,
        "asset_class": "residential_mortgage", "value": 0.008,
        "value_date": date(2024, 12, 31), "period_years": 5,
        "geography": "AU", "url": "https://example.com",
        "retrieval_date": date(2025, 3, 1),
        "quality_score": QualityScore.HIGH,
    }
    base.update(overrides)
    return BenchmarkEntry(**base)


@pytest.fixture()
def seeded_reporter():
    engine = create_engine_and_schema(":memory:")
    registry = BenchmarkRegistry(engine, actor="test")
    load_seed_data(registry)
    return GovernanceReporter(registry, InstitutionType.BANK), registry, engine


# ---------------------------------------------------------------------------
# Refresh schedules YAML loads with all 10 source types
# ---------------------------------------------------------------------------

def test_refresh_schedules_yaml_loads() -> None:
    schedules = load_refresh_schedules()
    required_keys = {
        "pillar3", "apra_adi", "rating_agency", "icc_trade",
        "industry_body", "listed_peer", "rba", "bureau",
        "insolvency", "regulatory",
    }
    assert required_keys.issubset(schedules.keys())
    assert schedules["pillar3"] == 120
    assert schedules["rating_agency"] == 395
    assert schedules["regulatory"] == 760


# ---------------------------------------------------------------------------
# 1. Stale check — per-source-type thresholds
# ---------------------------------------------------------------------------

def test_pillar3_at_150_days_is_stale_but_rating_agency_is_not() -> None:
    """Key Tier 5 test from the user brief — verify per-source-type thresholds."""
    engine = create_engine_and_schema(":memory:")
    registry = BenchmarkRegistry(engine, actor="test")
    # 150 days before 2025-08-01 = 2025-03-04
    retrieval_150d_ago = date(2025, 3, 4)
    as_of = date(2025, 8, 1)

    registry.add(_entry(
        source_id="CBA_PILLAR3", source_type=SourceType.PILLAR3,
        retrieval_date=retrieval_150d_ago,
    ))
    registry.add(_entry(
        source_id="SP_RATING", source_type=SourceType.RATING_AGENCY,
        retrieval_date=retrieval_150d_ago, value=0.017,
    ))

    reporter = GovernanceReporter(registry, InstitutionType.BANK)
    report = reporter.stale_benchmark_report(as_of=as_of)

    stale_ids = {f["source_id"] for f in report.findings if f["stale"]}
    assert "CBA_PILLAR3" in stale_ids           # 150 > 120 pillar3 threshold
    assert "SP_RATING" not in stale_ids         # 150 <= 395 rating_agency threshold
    assert any("stale:CBA_PILLAR3" in flag for flag in report.flags)
    assert not any("SP_RATING" in flag for flag in report.flags)


def test_stale_report_counts_all_entries(seeded_reporter) -> None:
    reporter, registry, _ = seeded_reporter
    seed_entries = registry.list()
    report = reporter.stale_benchmark_report()
    assert len(report.findings) == len(seed_entries)


# ---------------------------------------------------------------------------
# 2. Quality assessment — 5-dimension matrix
# ---------------------------------------------------------------------------

def test_quality_assessment_produces_five_dimension_scores(seeded_reporter) -> None:
    reporter, _, _ = seeded_reporter
    report = reporter.quality_assessment_report()

    assert report.report_type == "quality_assessment"
    assert len(report.findings) > 0
    for finding in report.findings:
        dims = finding["dimensions"]
        assert set(dims.keys()) == {
            "depth", "relevance", "transparency", "frequency", "regulatory_standing",
        }


def test_quality_assessment_flags_industry_body_low_regulatory_standing(
    seeded_reporter,
) -> None:
    reporter, _, _ = seeded_reporter
    report = reporter.quality_assessment_report()
    # CoreLogic, JLL, AFIA (industry_body) should have LOW regulatory_standing -> flag
    low_flags = [f for f in report.flags if f.startswith("low_quality:")]
    assert len(low_flags) > 0
    assert any("regulatory_standing" in f for f in low_flags)


# ---------------------------------------------------------------------------
# 3. Peer comparison — >30% divergence flagged
# ---------------------------------------------------------------------------

def test_peer_comparison_flags_30pct_divergence(seeded_reporter) -> None:
    reporter, _, _ = seeded_reporter
    # Seed residential Big 4 PDs ~0.8% median; own PD = 1.5% is >50% divergence
    own = {"residential_mortgage": 0.015}
    report = reporter.peer_comparison_report(
        own, segments=["residential_mortgage"],
    )
    finding = report.findings[0]
    assert finding["breached"] is True
    assert any("divergence:residential_mortgage" in f for f in report.flags)


def test_peer_comparison_within_tolerance_not_flagged(seeded_reporter) -> None:
    reporter, _, _ = seeded_reporter
    own = {"residential_mortgage": 0.0083}  # very close to peer median
    report = reporter.peer_comparison_report(own, segments=["residential_mortgage"])
    finding = report.findings[0]
    assert finding["breached"] is False
    assert report.flags == []


def test_peer_comparison_records_missing_own_estimate(seeded_reporter) -> None:
    reporter, _, _ = seeded_reporter
    report = reporter.peer_comparison_report(
        {}, segments=["residential_mortgage"],
    )
    assert report.findings[0]["status"] == "no_own_estimate"


# ---------------------------------------------------------------------------
# 4. Coverage report
# ---------------------------------------------------------------------------

def test_coverage_report_flags_sparse_segments(seeded_reporter) -> None:
    reporter, registry, _ = seeded_reporter
    # Add a segment with only one source
    registry.add(_entry(
        source_id="LONELY", asset_class="sparse_segment",
    ))
    report = reporter.coverage_report(
        segments=["residential_mortgage", "sparse_segment"],
    )
    sparse = next(f for f in report.findings if f["segment"] == "sparse_segment")
    dense = next(f for f in report.findings if f["segment"] == "residential_mortgage")
    assert sparse["meets_threshold"] is False
    assert dense["meets_threshold"] is True
    assert any("low_coverage:sparse_segment" in f for f in report.flags)


# ---------------------------------------------------------------------------
# 5. Annual review package — institution-specific committee header
# ---------------------------------------------------------------------------

def test_annual_review_bank_uses_mrc_header(seeded_reporter) -> None:
    reporter, _, _ = seeded_reporter
    report = reporter.annual_review_package(
        segments=["residential_mortgage", "commercial_property_investment"],
    )
    header = report.findings[0]
    assert header["section"] == "header"
    assert header["committee"] == "MRC"


def test_annual_review_pc_uses_credit_committee_header() -> None:
    engine = create_engine_and_schema(":memory:")
    registry = BenchmarkRegistry(engine, actor="test")
    load_seed_data(registry)
    reporter = GovernanceReporter(registry, InstitutionType.PRIVATE_CREDIT)
    report = reporter.annual_review_package(segments=["bridging_residential"])
    header = report.findings[0]
    assert header["committee"] == "Credit Committee"


def test_annual_review_includes_peer_when_own_estimates_given(seeded_reporter) -> None:
    reporter, _, _ = seeded_reporter
    report = reporter.annual_review_package(
        segments=["residential_mortgage"],
        own_estimates={"residential_mortgage": 0.0083},
    )
    sections = {f["section"] for f in report.findings}
    assert "peer_comparison" in sections


# ---------------------------------------------------------------------------
# 6. Version drift report
# ---------------------------------------------------------------------------

def test_version_drift_lists_all_versions_ascending() -> None:
    engine = create_engine_and_schema(":memory:")
    registry = BenchmarkRegistry(engine, actor="test")
    registry.add(_entry(source_id="S", value=0.08))
    registry.supersede("S", _entry(source_id="S", value=0.085))
    registry.supersede("S", _entry(source_id="S", value=0.092))

    reporter = GovernanceReporter(registry, InstitutionType.BANK)
    report = reporter.version_drift_report("S")

    versions = [f["version"] for f in report.findings]
    assert versions == [1, 2, 3]
    values = [f["value"] for f in report.findings]
    assert values == [0.08, 0.085, 0.092]


def test_version_drift_flags_large_change() -> None:
    engine = create_engine_and_schema(":memory:")
    registry = BenchmarkRegistry(engine, actor="test")
    registry.add(_entry(source_id="S", value=0.02))
    registry.supersede("S", _entry(source_id="S", value=0.05))  # +150%
    reporter = GovernanceReporter(registry, InstitutionType.BANK)
    report = reporter.version_drift_report("S")
    assert any("drift:S" in f for f in report.flags)


# ---------------------------------------------------------------------------
# Pillar 3 peer divergence (Part 1 of the reporting polish spec)
# ---------------------------------------------------------------------------

def _pillar3_entry(publisher: str, source_id: str, value: float, **overrides):
    """Helper: build a Pillar 3 residential-mortgage PD entry for peer tests."""
    base = _entry(
        source_id=source_id, publisher=publisher,
        source_type=SourceType.PILLAR3, data_type=DataType.PD,
        asset_class="residential_mortgage", value=value,
    )
    return base


def test_pillar3_peer_divergence_fires_on_outlier() -> None:
    engine = create_engine_and_schema(":memory:")
    registry = BenchmarkRegistry(engine, actor="test")
    # Three peers around 0.008, one outlier at 0.040 (5x median)
    registry.add(_pillar3_entry("CBA", "CBA_RES_PD", 0.0072))
    registry.add(_pillar3_entry("NAB", "NAB_RES_PD", 0.0090))
    registry.add(_pillar3_entry("WBC", "WBC_RES_PD", 0.0088))
    registry.add(_pillar3_entry("ANZ", "ANZ_RES_PD", 0.0400))  # outlier

    reporter = GovernanceReporter(registry, InstitutionType.BANK)
    report = reporter.pillar3_peer_divergence_report()

    assert report.report_type == "pillar3_peer_divergence"
    # Flag must name the outlier bank's source_id
    assert any("ANZ_RES_PD" in f for f in report.flags)
    # Non-outliers not flagged
    assert not any("CBA_RES_PD" in f for f in report.flags)


def test_pillar3_peer_divergence_quiet_on_normal_big4_range() -> None:
    """Natural Big 4 variation (0.0072–0.0090) must not flag anything."""
    engine = create_engine_and_schema(":memory:")
    registry = BenchmarkRegistry(engine, actor="test")
    for pub, sid, v in [
        ("CBA", "CBA_RES_PD", 0.0072),
        ("NAB", "NAB_RES_PD", 0.0090),
        ("WBC", "WBC_RES_PD", 0.0088),
        ("ANZ", "ANZ_RES_PD", 0.0080),
    ]:
        registry.add(_pillar3_entry(pub, sid, v))

    reporter = GovernanceReporter(registry, InstitutionType.BANK)
    report = reporter.pillar3_peer_divergence_report()
    assert report.flags == []
    # Every finding records a (breached=False) row
    assert all(f["breached"] is False for f in report.findings
               if f.get("status") != "skipped_incomplete_cohort")


def test_pillar3_peer_divergence_skips_incomplete_cohorts() -> None:
    """Fewer than min_cohort_size banks -> skip with an explanatory row."""
    engine = create_engine_and_schema(":memory:")
    registry = BenchmarkRegistry(engine, actor="test")
    registry.add(_pillar3_entry("CBA", "CBA_RES_PD", 0.0072))
    registry.add(_pillar3_entry("NAB", "NAB_RES_PD", 0.0400))  # would be outlier if evaluated

    reporter = GovernanceReporter(registry, InstitutionType.BANK)
    report = reporter.pillar3_peer_divergence_report(min_cohort_size=3)

    # No flag because cohort size 2 < 3
    assert report.flags == []
    skipped = [f for f in report.findings
               if f.get("status") == "skipped_incomplete_cohort"]
    assert len(skipped) == 1
    assert skipped[0]["cohort_size"] == 2


def test_annual_review_package_includes_pillar3_peer_divergence(seeded_reporter) -> None:
    """New section appears whenever annual_review_package() runs."""
    reporter, _, _ = seeded_reporter
    report = reporter.annual_review_package(
        segments=["residential_mortgage"],
    )
    sections = {f["section"] for f in report.findings}
    assert "pillar3_peer_divergence" in sections


# ---------------------------------------------------------------------------
# READ-ONLY enforcement: governance must not write to DB
# ---------------------------------------------------------------------------

def test_governance_does_not_write_to_adjustments_or_audit_log(seeded_reporter) -> None:
    reporter, _, engine = seeded_reporter
    factory = make_session_factory(engine)

    def counts():
        with factory() as s:
            adj = s.scalars(select(func.count(Adjustment.pk))).one()
            audit_adjust = s.scalars(
                select(func.count(AuditLog.pk)).where(AuditLog.operation == "adjust")
            ).one()
        return adj, audit_adjust

    before = counts()
    reporter.stale_benchmark_report()
    reporter.quality_assessment_report()
    reporter.coverage_report(segments=["residential_mortgage"])
    reporter.peer_comparison_report({"residential_mortgage": 0.008}, ["residential_mortgage"])
    reporter.annual_review_package(segments=["residential_mortgage"])
    after = counts()

    # Registry reads (list/get_by_segment/get_version_history) DO write audit_log —
    # but `operation='adjust'` rows should never appear from governance.
    assert before == after


# ---------------------------------------------------------------------------
# DOCX export — ImportError shim when python-docx missing
# ---------------------------------------------------------------------------

def test_export_to_docx_raises_with_install_hint_when_docx_missing(
    seeded_reporter, tmp_path
) -> None:
    reporter, _, _ = seeded_reporter
    report = reporter.stale_benchmark_report()

    # Simulate the dependency being absent by making the import fail.
    import builtins
    real_import = builtins.__import__

    def blocked_import(name, *args, **kwargs):
        if name == "docx" or name.startswith("docx."):
            raise ImportError("No module named 'docx'")
        return real_import(name, *args, **kwargs)

    with patch.object(builtins, "__import__", side_effect=blocked_import):
        with pytest.raises(ImportError, match=r"pip install.*\[reports\]"):
            export_to_docx(report, tmp_path / "report.docx")


# ---------------------------------------------------------------------------
# Part 2: DOCX polish — MRC (bank) vs Credit Committee (PC) formatting
# ---------------------------------------------------------------------------

def _reopen_docx_text(path) -> str:
    """Open a DOCX file and return all paragraph + table cell text concatenated."""
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def test_export_to_docx_bank_produces_non_empty_file(
    seeded_reporter, tmp_path,
) -> None:
    reporter, _, _ = seeded_reporter
    report = reporter.stale_benchmark_report()
    path = tmp_path / "bank.docx"
    export_to_docx(report, path, institution_type="bank")
    assert path.exists()
    assert path.stat().st_size > 1000   # DOCX minimum plausible size


def test_export_to_docx_pc_produces_non_empty_file(
    seeded_reporter, tmp_path,
) -> None:
    reporter, _, _ = seeded_reporter
    report = reporter.stale_benchmark_report()
    path = tmp_path / "pc.docx"
    export_to_docx(report, path, institution_type="private_credit")
    assert path.exists()
    assert path.stat().st_size > 1000


def test_export_to_docx_bank_includes_3lod_and_model_owner(
    seeded_reporter, tmp_path,
) -> None:
    reporter, _, _ = seeded_reporter
    report = reporter.stale_benchmark_report()
    path = tmp_path / "bank.docx"
    export_to_docx(report, path, institution_type="bank")

    text = _reopen_docx_text(path)
    assert "Model Risk Committee" in text
    assert "3 Lines of Defence Sign-Off" in text
    assert "Model Owner" in text
    assert "Model Validation" in text
    assert "Internal Audit" in text


def test_export_to_docx_pc_includes_decision_log_and_next_actions(
    seeded_reporter, tmp_path,
) -> None:
    reporter, _, _ = seeded_reporter
    report = reporter.stale_benchmark_report()
    path = tmp_path / "pc.docx"
    export_to_docx(report, path, institution_type="private_credit")

    text = _reopen_docx_text(path)
    assert "Credit Committee" in text
    assert "Decision Log" in text
    assert "Next Review Actions" in text


def test_export_to_docx_defaults_to_report_institution_when_type_omitted(
    tmp_path,
) -> None:
    """When no institution_type kwarg, fall back to report.institution_type."""
    engine = create_engine_and_schema(":memory:")
    registry = BenchmarkRegistry(engine, actor="test")
    reporter = GovernanceReporter(registry, InstitutionType.PRIVATE_CREDIT)
    report = reporter.stale_benchmark_report()

    path = tmp_path / "auto.docx"
    export_to_docx(report, path)   # no institution_type
    text = _reopen_docx_text(path)
    assert "Credit Committee" in text
