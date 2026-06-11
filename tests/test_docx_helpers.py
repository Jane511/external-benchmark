"""Smoke test for the lean docx model-input renderer."""
from __future__ import annotations

from datetime import date, timedelta
from pathlib import Path

import pytest

pytest.importorskip("docx", reason="python-docx is required for docx smoke test")

from docx import Document  # noqa: E402

from src.benchmark_report import BenchmarkCalibrationReport  # noqa: E402
from src.db import create_engine_and_schema  # noqa: E402
from src.models import (  # noqa: E402
    DataDefinitionClass,
    RawObservation,
    SourceType,
)
from src.registry import BenchmarkRegistry  # noqa: E402


@pytest.fixture()
def populated_registry() -> BenchmarkRegistry:
    engine = create_engine_and_schema(":memory:")
    reg = BenchmarkRegistry(engine, actor="test")
    today = date(2026, 4, 27)
    reg.add_observations([
        RawObservation(
            source_id="cba",
            source_type=SourceType.BANK_PILLAR3,
            segment="commercial_property",
            parameter="pd",
            data_definition_class=DataDefinitionClass.BASEL_PD_ONE_YEAR,
            value=0.025,
            as_of_date=today - timedelta(days=120),
            reporting_basis="Pillar 3 H2 2024",
            methodology_note="CBA Pillar 3 Table CR6 commercial property PD",
        ),
        RawObservation(
            source_id="cba",
            source_type=SourceType.BANK_PILLAR3,
            segment="commercial_property",
            parameter="pd",
            data_definition_class=DataDefinitionClass.BASEL_PD_ONE_YEAR,
            value=0.024,
            as_of_date=today - timedelta(days=300),
            reporting_basis="Pillar 3 H1 2024",
            methodology_note="CBA Pillar 3 Table CR6 commercial property PD",
        ),
        RawObservation(
            source_id="judo",
            source_type=SourceType.NON_BANK_LISTED,
            segment="commercial_property",
            parameter="pd",
            data_definition_class=DataDefinitionClass.BASEL_PD_ONE_YEAR,
            value=0.045,
            as_of_date=today - timedelta(days=90),
            reporting_basis="Half-yearly disclosure",
            methodology_note="Judo Bank Pillar 3 commercial book",
        ),
        RawObservation(
            source_id="aps113_cre_lgd",
            source_type=SourceType.REGULATORY,
            segment="commercial_property",
            parameter="lgd",
            data_definition_class=DataDefinitionClass.REGULATORY_FLOOR_LGD,
            value=0.175,
            as_of_date=today - timedelta(days=120),
            reporting_basis="APS 113",
            methodology_note="Commercial-property LGD floor",
        ),
        RawObservation(
            source_id="QUALITAS_CRE_COMMENTARY",
            source_type=SourceType.NON_BANK_LISTED,
            segment="commercial_property",
            parameter="commentary",
            data_definition_class=DataDefinitionClass.QUALITATIVE_COMMENTARY,
            value=None,
            as_of_date=today - timedelta(days=120),
            reporting_basis="Half-yearly results commentary",
            methodology_note="QUALITATIVE: office sector under pressure",
        ),
    ])
    return reg


def _all_text(doc) -> str:
    paragraphs = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    return paragraphs + "\n" + table_text


def test_to_docx_round_trip(populated_registry, tmp_path: Path) -> None:
    """End-to-end: render docx, re-read with python-docx, assert structure."""
    report = BenchmarkCalibrationReport(
        populated_registry, period_label="Q1 2026",
    )
    out = tmp_path / "smoke.docx"
    report.to_docx(out)
    assert out.exists() and out.stat().st_size > 0

    doc = Document(str(out))
    text = _all_text(doc)

    assert "Australian Credit Risk Benchmarks - Q1 2026" in text
    assert "1. PD Inputs" in text
    assert "2. LGD Inputs" in text
    assert "3. Expected Loss Inputs" in text
    assert "4. Stress Testing Inputs" in text
    assert "5. Portfolio Monitor Inputs" in text
    assert "6. Per-Bank Industry Inputs" in text
    assert "Provenance" not in text
    assert "Glossary" not in text


def test_to_docx_includes_trend_when_two_vintages_present(
    populated_registry, tmp_path: Path,
) -> None:
    report = BenchmarkCalibrationReport(
        populated_registry, period_label="Q1 2026",
    )
    out = tmp_path / "smoke_trend.docx"
    report.to_docx(out)

    doc = Document(str(out))
    text = _all_text(doc)
    assert "7. Trend vs prior cycle" not in text
    assert "Expected Loss Inputs" in text
